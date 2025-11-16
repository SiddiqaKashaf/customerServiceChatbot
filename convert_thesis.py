"""
Script to convert THESIS.md to DOCX format
"""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_run_font(run, font_name='Times New Roman', font_size=12, bold=False, italic=False):
    """Set font properties for a run"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    # Set Chinese font support
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_heading(doc, text, level=1):
    """Add a heading to the document"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        set_run_font(run, font_name='Times New Roman', font_size=14 + (3-level)*2, bold=True)
    return heading

def add_paragraph(doc, text, style='Normal', bold=False, italic=False):
    """Add a paragraph to the document"""
    para = doc.add_paragraph(style=style)
    run = para.add_run(text)
    set_run_font(run, font_name='Times New Roman', font_size=12, bold=bold, italic=italic)
    return para

def parse_markdown_to_docx(md_file, docx_file):
    """Convert Markdown file to DOCX format"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Title (single #)
        if line.startswith('# ') and not line.startswith('##'):
            title = line[2:].strip()
            # Skip if it's just the main title (already in header)
            if 'AI-Powered' in title or 'Final Year Project' in title:
                i += 1
                continue
            add_heading(doc, title, level=1)
        
        # Chapter heading
        elif line.startswith('# Chapter'):
            title = line[2:].strip()
            add_heading(doc, title, level=1)
        
        # Section heading (##)
        elif line.startswith('## ') and not line.startswith('###'):
            title = line[3:].strip()
            if title.startswith('Chapter'):
                add_heading(doc, title, level=1)
            else:
                add_heading(doc, title, level=2)
        
        # Subsection heading (###)
        elif line.startswith('### '):
            title = line[4:].strip()
            add_heading(doc, title, level=3)
        
        # Subsubsection heading (####)
        elif line.startswith('#### '):
            title = line[5:].strip()
            add_heading(doc, title, level=4)
        
        # Horizontal rule
        elif line.startswith('---'):
            # Add some space
            doc.add_paragraph()
        
        # Bold text
        elif '**' in line:
            para = doc.add_paragraph()
            # Simple bold handling
            parts = re.split(r'(\*\*[^*]+\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = para.add_run(part[2:-2])
                    set_run_font(run, bold=True)
                elif part:
                    run = para.add_run(part)
                    set_run_font(run)
        
        # Bullet points
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            # Remove bold markers for now
            text = text.replace('**', '')
            para = doc.add_paragraph(text, style='List Bullet')
            for run in para.runs:
                set_run_font(run)
        
        # Numbered list
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            text = text.replace('**', '')
            para = doc.add_paragraph(text, style='List Number')
            for run in para.runs:
                set_run_font(run)
        
        # Code blocks (skip for now or format differently)
        elif line.startswith('```'):
            # Skip code blocks
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                i += 1
        
        # Regular paragraph
        else:
            # Remove markdown formatting
            text = line
            text = text.replace('**', '')
            text = text.replace('*', '')
            # Remove links [text](url) -> text
            text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
            
            if text.strip():
                para = doc.add_paragraph(text)
                for run in para.runs:
                    set_run_font(run)
        
        i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"Successfully converted {md_file} to {docx_file}")

if __name__ == '__main__':
    try:
        parse_markdown_to_docx('THESIS.md', 'THESIS.docx')
        print("Conversion completed successfully!")
    except Exception as e:
        print(f"Error during conversion: {e}")
        import traceback
        traceback.print_exc()

