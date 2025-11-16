"""
Advanced script to convert THESIS.md to DOCX and PDF formats
"""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_run_font(run, font_name='Times New Roman', font_size=12, bold=False, italic=False):
    """Set font properties for a run"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    # Set Chinese font support
    try:
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    except:
        pass

def add_heading(doc, text, level=1):
    """Add a heading to the document"""
    # Clean text
    text = text.strip()
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)  # Remove markdown links
    
    heading = doc.add_heading(text, level=min(level, 9))
    for run in heading.runs:
        font_size = 16 if level == 1 else (14 if level == 2 else (12 if level == 3 else 11))
        set_run_font(run, font_name='Times New Roman', font_size=font_size, bold=True)
    return heading

def add_formatted_paragraph(doc, text, style='Normal'):
    """Add a paragraph with markdown formatting"""
    para = doc.add_paragraph(style=style)
    
    # Handle bold text
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            set_run_font(run, bold=True)
        elif part:
            run = para.add_run(part)
            set_run_font(run)
    
    return para

def parse_markdown_to_docx(md_file, docx_file):
    """Convert Markdown file to DOCX format with better formatting"""
    doc = Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    
    i = 0
    in_code_block = False
    code_block_content = []
    
    while i < len(lines):
        line = lines[i]
        original_line = line
        line = line.rstrip()
        
        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End of code block
                if code_block_content:
                    para = doc.add_paragraph('\n'.join(code_block_content), style='Normal')
                    para_format = para.paragraph_format
                    para_format.left_indent = Inches(0.5)
                    for run in para.runs:
                        set_run_font(run, font_name='Courier New', font_size=10)
                code_block_content = []
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_block_content.append(line)
            i += 1
            continue
        
        # Skip empty lines (but add spacing between sections)
        if not line.strip():
            # Add paragraph for spacing
            doc.add_paragraph()
            i += 1
            continue
        
        # Title page content
        if i < 10 and ('Author:' in line or 'Supervisor:' in line or 'Institution:' in line or 'Date:' in line):
            para = doc.add_paragraph()
            if '**' in line:
                parts = re.split(r'(\*\*[^*]+\*\*:)', line)
                for part in parts:
                    if part.endswith(':'):
                        run = para.add_run(part)
                        set_run_font(run, bold=True)
                    elif part and not part.startswith('**'):
                        run = para.add_run(part)
                        set_run_font(run)
            else:
                run = para.add_run(line)
                set_run_font(run)
            i += 1
            continue
        
        # Main title
        if line.startswith('# ') and 'AI-Powered' in line:
            title = line[2:].strip()
            para = doc.add_heading(title, level=1)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                set_run_font(run, font_size=18, bold=True)
            i += 1
            continue
        
        # Chapter heading
        if line.startswith('# Chapter'):
            title = line[2:].strip()
            # Add page break before new chapter (except first)
            if i > 50:  # Not the first chapter
                doc.add_page_break()
            add_heading(doc, title, level=1)
            i += 1
            continue
        
        # Section heading (##)
        if line.startswith('## ') and not line.startswith('###'):
            title = line[3:].strip()
            # Remove markdown links
            title = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', title)
            add_heading(doc, title, level=2)
            i += 1
            continue
        
        # Subsection heading (###)
        if line.startswith('### '):
            title = line[4:].strip()
            title = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', title)
            add_heading(doc, title, level=3)
            i += 1
            continue
        
        # Subsubsection heading (####)
        if line.startswith('#### '):
            title = line[5:].strip()
            title = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', title)
            add_heading(doc, title, level=4)
            i += 1
            continue
        
        # Horizontal rule
        if line.startswith('---'):
            doc.add_paragraph()
            i += 1
            continue
        
        # Bullet points
        if line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            # Handle nested lists
            indent_level = 0
            while text.startswith('  '):
                text = text[2:]
                indent_level += 1
            
            # Remove bold markers but preserve structure
            para = doc.add_paragraph(text, style='List Bullet')
            para_format = para.paragraph_format
            para_format.left_indent = Inches(0.25 * (indent_level + 1))
            
            # Format bold text in list items
            if '**' in text:
                para.clear()
                parts = re.split(r'(\*\*[^*]+\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = para.add_run(part[2:-2])
                        set_run_font(run, bold=True)
                    elif part:
                        run = para.add_run(part)
                        set_run_font(run)
            else:
                for run in para.runs:
                    set_run_font(run)
            i += 1
            continue
        
        # Numbered list
        match = re.match(r'^(\d+)\.\s(.+)$', line)
        if match:
            text = match.group(2).strip()
            para = doc.add_paragraph(text, style='List Number')
            
            # Format bold text in numbered items
            if '**' in text:
                para.clear()
                parts = re.split(r'(\*\*[^*]+\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = para.add_run(part[2:-2])
                        set_run_font(run, bold=True)
                    elif part:
                        run = para.add_run(part)
                        set_run_font(run)
            else:
                for run in para.runs:
                    set_run_font(run)
            i += 1
            continue
        
        # Regular paragraph with formatting
        if line.strip():
            # Remove markdown links but keep text
            text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', line)
            
            # Handle paragraphs with bold
            if '**' in text:
                add_formatted_paragraph(doc, text)
            else:
                para = doc.add_paragraph(text)
                for run in para.runs:
                    set_run_font(run)
        
        i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"✓ Successfully converted {md_file} to {docx_file}")

def create_pdf_from_docx(docx_file, pdf_file):
    """Convert DOCX to PDF using docx2pdf (requires Microsoft Word on Windows)"""
    try:
        from docx2pdf import convert
        convert(docx_file, pdf_file)
        print(f"✓ Successfully converted {docx_file} to {pdf_file}")
        return True
    except ImportError:
        print("⚠ docx2pdf not available. Installing...")
        try:
            import subprocess
            subprocess.check_call(['pip', 'install', 'docx2pdf'])
            from docx2pdf import convert
            convert(docx_file, pdf_file)
            print(f"✓ Successfully converted {docx_file} to {pdf_file}")
            return True
        except Exception as e:
            print(f"✗ Could not convert to PDF: {e}")
            print("  Note: PDF conversion requires Microsoft Word on Windows")
            print("  Alternative: Open THESIS.docx in Word and save as PDF")
            return False
    except Exception as e:
        print(f"✗ Could not convert to PDF: {e}")
        print("  Note: PDF conversion requires Microsoft Word on Windows")
        print("  Alternative: Open THESIS.docx in Word and save as PDF")
        return False

if __name__ == '__main__':
    import sys
    import os
    
    # Use different output filename to avoid permission issues
    docx_output = 'THESIS_Final.docx'
    pdf_output = 'THESIS_Final.pdf'
    
    print("Converting THESIS.md to DOCX...")
    try:
        # Remove existing file if it exists
        if os.path.exists(docx_output):
            try:
                os.remove(docx_output)
            except:
                pass
        
        parse_markdown_to_docx('THESIS.md', docx_output)
        print(f"\n✓ DOCX conversion completed successfully! Saved as {docx_output}")
        
        # Try to create PDF
        print("\nAttempting to create PDF...")
        if create_pdf_from_docx(docx_output, pdf_output):
            print(f"✓ PDF conversion completed successfully! Saved as {pdf_output}")
        
    except Exception as e:
        print(f"✗ Error during conversion: {e}")
        import traceback
        traceback.print_exc()
        print("\nNote: If you have the DOCX file open, please close it and try again.")

